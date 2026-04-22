"""
NOC Report Generator
Produces a .docx report matching the sample format:
  - Header: Date / Shift / Agent
  - Section 1: System Incidents
  - Section 2: Fibre Incidents (grouped by region)
  - Section 3: System Uptime table
  - Screenshots embedded at end
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.database import (
    get_shift, get_system_incidents, get_fibre_incidents,
    get_uptime, get_screenshots, get_regional_incidents
)

REPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "reports")
YELLOW = RGBColor(0xFF, 0xFF, 0x00)
ORANGE = RGBColor(0xFF, 0xC0, 0x00)
BLUE_HDR = RGBColor(0xBD, 0xD7, 0xEE)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x40, 0x40, 0x40)


def _set_cell_bg(cell, hex_color: str):
    """Set cell background shading via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, top=True, bottom=True, left=True, right=True):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, on in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single" if on else "none")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_handover_notes(cell, notes_text: str, size=8):
    """
    Render handover notes with bold timestamps per line.
    Each entry on its own paragraph, blank line between entries.
    """
    import re
    TIMESTAMP_RE = re.compile(
        r'^(\d{1,2}/\d{1,2}/\d{2,4}\s*[@]\s*\d{3,4}[Hh][Rr][Ss]?\s*(?:–|-|:)?)\s*(.*)',
        re.DOTALL
    )
    lines = (notes_text or "").strip().split("\n")
    first_para = cell.paragraphs[0]
    # clear default empty paragraph
    first_para.clear()
    first_para.paragraph_format.space_after = Pt(0)

    used_first = False
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        if used_first:
            cell.add_paragraph()  # blank spacer between entries
            para = cell.add_paragraph()
        else:
            para = first_para
            used_first = True
        para.paragraph_format.space_after = Pt(0)
        m = TIMESTAMP_RE.match(line)
        if m:
            r1 = para.add_run(m.group(1))
            r1.bold = True
            r1.font.size = Pt(size)
            rest = m.group(2).strip()
            if rest:
                r2 = para.add_run(" " + rest)
                r2.bold = False
                r2.font.size = Pt(size)
        else:
            r = para.add_run(line)
            r.font.size = Pt(size)


def _para(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para


def _add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK


def generate_report(shift_id: int) -> str:
    """Generate the full NOC shift report. Returns the file path."""
    os.makedirs(REPORTS_DIR, exist_ok=True)

    shift = get_shift(shift_id)
    if not shift:
        raise ValueError(f"Shift {shift_id} not found")

    sys_incidents = get_system_incidents(shift_id)
    regional_incidents = get_regional_incidents(shift_id)
    fibre_incidents = get_fibre_incidents(shift_id)
    uptime_data = get_uptime(shift_id)
    screenshots = get_screenshots(shift_id)

    doc = Document()

    # ── Page margins ─────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Default style ────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # HEADER TABLE: Date | Shift | Agent
    hdr_table = doc.add_table(rows=1, cols=3)
    hdr_table.style = "Table Grid"
    hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    date_str = shift["start_time"][:10] if shift["start_time"] else datetime.now().strftime("%d/%m/%Y")
    try:
        dt = datetime.strptime(shift["start_time"][:10], "%Y-%m-%d")
        date_str = dt.strftime("%d/%m/%Y")
    except Exception:
        pass

    hdr_cells = hdr_table.rows[0].cells
    for cell in hdr_cells:
        _set_cell_bg(cell, "F2F2F2")

    # Date cell
    hdr_cells[0].paragraphs[0].clear()
    p = hdr_cells[0].paragraphs[0]
    r1 = p.add_run("Date: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    r1.font.size = Pt(9)
    r2 = p.add_run(date_str)
    r2.bold = True
    r2.font.size = Pt(9)

    # Shift cell
    hdr_cells[1].paragraphs[0].clear()
    p = hdr_cells[1].paragraphs[0]
    r1 = p.add_run("Shift: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    r1.font.size = Pt(9)
    r2 = p.add_run(shift["shift_type"])
    r2.bold = True
    r2.font.size = Pt(9)

    # Agent cell
    hdr_cells[2].paragraphs[0].clear()
    p = hdr_cells[2].paragraphs[0]
    r1 = p.add_run("Agent: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    r1.font.size = Pt(9)
    r2 = p.add_run(shift["agent_name"])
    r2.bold = True
    r2.font.size = Pt(9)

    doc.add_paragraph()

    # SECTION 1 — SYSTEM INCIDENTS
    _add_section_heading(doc, "Incidences on critical systems/service")

    sys_cols = ["Description", "Date/Time", "End Date/\nTime", "Incident\nNo",
                "Duration", "Action to", "Status",
                "Incident activities and handover notes", "Incident Report\nprovided"]
    sys_widths = [Cm(3.2), Cm(2.8), Cm(2.0), Cm(1.6), Cm(1.6), Cm(1.5),
                  Cm(1.8), Cm(6.5), Cm(2.0)]

    sys_table = doc.add_table(rows=1, cols=len(sys_cols))
    sys_table.style = "Table Grid"

    # Header row
    hrow = sys_table.rows[0]
    for i, (col, w) in enumerate(zip(sys_cols, sys_widths)):
        cell = hrow.cells[i]
        cell.width = w
        _set_cell_bg(cell, "FFFF00")
        _set_cell_borders(cell)
        cell.paragraphs[0].clear()
        _para(cell, col, bold=True, size=8)

    # Data rows
    if not sys_incidents:
        row = sys_table.add_row()
        row.cells[0].paragraphs[0].clear()
        p = row.cells[0].paragraphs[0]
        run = p.add_run("No system incidents this shift.")
        run.font.size = Pt(8)
        row.cells[0].merge(row.cells[-1])
    else:
        for inc in sys_incidents:
            row = sys_table.add_row()
            cells = row.cells
            vals = [
                inc.get("description", ""),
                inc.get("date_time", ""),
                inc.get("end_time", ""),
                inc.get("incident_no", ""),
                inc.get("duration", ""),
                inc.get("action_to", ""),
                inc.get("status", ""),
                inc.get("activities", ""),
                "✓" if inc.get("report_provided") else "",
            ]
            for i, (val, w) in enumerate(zip(vals, sys_widths)):
                cells[i].width = w
                _set_cell_bg(cells[i], "FFFFFF")
                _set_cell_borders(cells[i])
                if i == 7:  # "Incident activities and handover notes"
                    _para_handover_notes(cells[i], val, size=8)
                else:
                    cells[i].paragraphs[0].clear()
                    bold_status = i == 6 and val in ("Resolved", "Assigned", "Ongoing", "Escalated")
                    _para(cells[i], val, bold=bold_status, size=8)

    doc.add_paragraph()

    # SECTION 2 — REGIONAL NETWORKS
    _add_section_heading(doc, "REGIONAL NETWORKS")

    reg_cols = ["Incidences", "Ref. No", "Date & Time\nReported", "Duration",
                "Person Assigned", "Report\nStatus", "Handover Notes"]
    reg_widths = [Cm(4.5), Cm(1.8), Cm(2.5), Cm(1.5), Cm(2.5), Cm(2.0), Cm(8.2)]

    reg_table = doc.add_table(rows=1, cols=len(reg_cols))
    reg_table.style = "Table Grid"

    # Header row
    rhrow = reg_table.rows[0]
    for i, (col, w) in enumerate(zip(reg_cols, reg_widths)):
        cell = rhrow.cells[i]
        cell.width = w
        _set_cell_bg(cell, "FFFF00")
        _set_cell_borders(cell)
        cell.paragraphs[0].clear()
        _para(cell, col, bold=True, size=8)

    if not regional_incidents:
        row = reg_table.add_row()
        row.cells[0].paragraphs[0].clear()
        run = row.cells[0].paragraphs[0].add_run("No regional network incidents this shift.")
        run.font.size = Pt(8)
        row.cells[0].merge(row.cells[-1])
    else:
        for inc in regional_incidents:
            row = reg_table.add_row()
            cells = row.cells
            vals = [
                inc.get("description", ""),
                inc.get("ref_no", ""),
                inc.get("date_reported", ""),
                inc.get("duration", ""),
                inc.get("person_assigned", ""),
                inc.get("report_status", ""),
                inc.get("handover_notes", ""),
            ]
            for i, (val, w) in enumerate(zip(vals, reg_widths)):
                cells[i].width = w
                _set_cell_bg(cells[i], "D9D9D9" if i == 0 else "FFFFFF")
                _set_cell_borders(cells[i])
                if i == 6:  # Handover Notes
                    _para_handover_notes(cells[i], val, size=8)
                else:
                    cells[i].paragraphs[0].clear()
                    bold = i == 0 or (i == 5 and val in ("Resolved", "Assigned", "Ongoing"))
                    _para(cells[i], val, bold=bold, size=8)

    doc.add_paragraph()

    # SECTION 3 — FIBRE INCIDENTS
    _add_section_heading(doc, "FIBER INCIDENCES")

    fibre_cols = ["Incidences", "Ref. No", "Date & Time\nReported", "Duration",
                  "Person Assigned", "Report Status", "Handover Notes"]
    fibre_widths = [Cm(4.5), Cm(1.8), Cm(2.5), Cm(1.5), Cm(2.5), Cm(2.0), Cm(8.2)]

    fibre_table = doc.add_table(rows=1, cols=len(fibre_cols))
    fibre_table.style = "Table Grid"

    # Header
    fhrow = fibre_table.rows[0]
    for i, (col, w) in enumerate(zip(fibre_cols, fibre_widths)):
        cell = fhrow.cells[i]
        cell.width = w
        _set_cell_bg(cell, "FFFF00")
        _set_cell_borders(cell)
        cell.paragraphs[0].clear()
        _para(cell, col, bold=True, size=8)

    # Group by region
    regions = {}
    for fi in fibre_incidents:
        r = fi.get("region", "NAIROBI REGION")
        regions.setdefault(r, []).append(fi)

    if not fibre_incidents:
        row = fibre_table.add_row()
        row.cells[0].paragraphs[0].clear()
        p = row.cells[0].paragraphs[0]
        run = p.add_run("No fibre incidents this shift.")
        run.font.size = Pt(8)
        row.cells[0].merge(row.cells[-1])
    else:
        for region, incidents in regions.items():
            # Region sub-header
            rrow = fibre_table.add_row()
            rc = rrow.cells[0]
            rc.merge(rrow.cells[-1])
            _set_cell_bg(rc, "FFC000")
            rc.paragraphs[0].clear()
            _para(rc, region, bold=False, size=8, color=RGBColor(0x7F, 0x3F, 0x00))

            for fi in incidents:
                row = fibre_table.add_row()
                cells = row.cells
                vals = [
                    fi.get("description", ""),
                    fi.get("ref_no", ""),
                    fi.get("date_reported", ""),
                    fi.get("duration", ""),
                    fi.get("person_assigned", ""),
                    fi.get("report_status", ""),
                    fi.get("handover_notes", ""),
                ]
                for i, (val, w) in enumerate(zip(vals, fibre_widths)):
                    cells[i].width = w
                    _set_cell_bg(cells[i], "D9D9D9" if i == 0 else "FFFFFF")
                    _set_cell_borders(cells[i])
                    if i == 6:  # Handover Notes column
                        _para_handover_notes(cells[i], val, size=8)
                    else:
                        cells[i].paragraphs[0].clear()
                        bold = i == 0 or (i == 5 and val in ("Resolved", "Assigned", "Ongoing"))
                        _para(cells[i], val, bold=bold, size=8)

    doc.add_paragraph()

    # SECTION 3 — SYSTEM UPTIME
    _add_section_heading(doc, "SYSTEM UPTIME (DAYS)")

    up_table = doc.add_table(rows=1, cols=3)
    up_table.style = "Table Grid"
    up_cols = ["SYSTEM", "UPTIME (DAYS)", "LAST DATE OF OUTAGE"]
    up_widths = [Cm(6), Cm(5), Cm(6)]

    uphrow = up_table.rows[0]
    for i, (col, w) in enumerate(zip(up_cols, up_widths)):
        cell = uphrow.cells[i]
        cell.width = w
        _set_cell_bg(cell, "BDD7EE")
        cell.paragraphs[0].clear()
        _para(cell, col, bold=True, size=9)

    if not uptime_data:
        # Default systems if nothing entered
        defaults = ["AMI", "INCMS", "INSMM", "USSD", "CONTACT CENTRE",
                    "PREPAID", "POSTPAID", "TOKEN TRACKER", "IPMP", "SAP"]
        for sys in defaults:
            row = up_table.add_row()
            row.cells[0].paragraphs[0].clear()
            _para(row.cells[0], sys, size=9)
            _para(row.cells[1], "", size=9)
            _para(row.cells[2], "", size=9)
    else:
        for item in uptime_data:
            row = up_table.add_row()
            row.cells[0].paragraphs[0].clear()
            _para(row.cells[0], item.get("system_name", ""), size=9)
            _para(row.cells[1], str(item.get("uptime_days", "") or "") + " days", size=9)
            _para(row.cells[2], item.get("last_outage_date", ""), size=9)

    doc.add_paragraph()

    # SECTION 4 — SCREENSHOTS
    if screenshots:
        _add_section_heading(doc, "SYSTEM SCREENSHOTS")
        for ss in screenshots:
            fp = ss.get("filepath", "")
            if fp and os.path.exists(fp):
                caption = ss.get("caption") or ss.get("system_name") or os.path.basename(fp)
                p = doc.add_paragraph(caption)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)
                try:
                    doc.add_picture(fp, width=Inches(9))
                except Exception:
                    p2 = doc.add_paragraph(f"[Screenshot: {fp}]")
                    p2.runs[0].font.size = Pt(8)
                doc.add_paragraph()

    # Save
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    agent_safe = "".join(c for c in shift["agent_name"] if c.isalnum() or c in " _-").strip()
    filename = f"NOC_Report_{agent_safe}_{ts}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)
    doc.save(filepath)
    print(f"[Report] Saved: {filepath}")
    return filepath
